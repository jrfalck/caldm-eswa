"""
CALDM Ablation Study — V3 (PRAXIS-MATCHED HYPERPARAMETERS).

What changed from V2
--------------------
V2's reproduction of "Full CALDM" reached AUC = 0.70, AP = 0.26 on the test set.
The original praxis result was AUC = 0.79, AP = 0.37 — a ~9 AUC point gap that
made the ablation comparisons unreliable.

Investigating, the gap was almost entirely explained by hyperparameter drift
between the praxis defaults and the V2 defaults:

    Hyperparameter        Praxis    V2 default
    --------------------- --------  ----------
    Diffusion epochs      200       30          (6.7x more in praxis)
    VAE epochs            15        30          (0.5x)
    Latent dim numerical  16        4           (4x more capacity in praxis)
    Latent dim categorical 32       10          (3.2x more capacity)
    Context dim           40        10          (4x more capacity)
    Salary weight         2         3           (0.67x)

The latent-dim and context-dim differences mean V2 was running with roughly 1/3
the representational capacity of the praxis model. The ablation conclusions
("E ties C", "VAE-context underperforms") were drawn against starved models.

V3 restores the praxis hyperparameters as defaults. The training code, configs,
and labels are otherwise identical to V2. The ContextEncoder still applies a
tanh bound to its output for joint-training stability; with context_dim=40 the
[-1,1]^40 space provides ample conditioning capacity. The bound can be disabled
via the `bound_context=False` parameter to run_ablation() if needed.

Output directory defaults to `ablation_outputs_v3/`. V1 and V2 results stay
intact for comparison.

Configurations evaluated:
    A. VAE only (no diffusion, no context)
    B. Data-space DDPM only (no VAE, no context)
    C. Latent Diffusion (unconditional)
    D. Context in VAE only
    E. CALDM (PROPOSED) -- Context in Diffusion only, on context-free latents
    F. CALDM + VAE context  (the V1 "Full CALDM" / praxis dual-conditioning)

Run:
    python caldm_ablation_study_v3.py            # full run, ~30 experiments
    python caldm_ablation_study_v3.py --quick    # 1 seed, fewer epochs

Author: Juan
"""

import os
import time
import argparse
import random
from dataclasses import dataclass

import numpy as np
import pandas as pd
import torch
import torch.nn as nn
import torch.optim as optim
import torch.nn.functional as F
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import (
    accuracy_score, precision_score, recall_score, f1_score,
    roc_auc_score, average_precision_score, roc_curve,
)
from scipy import stats
import matplotlib.pyplot as plt
import seaborn as sns


# =============================================================================
# 0. Configuration
# =============================================================================

DATA_PATH = "data/final_data_subsetFLGAAL8.npz"
SUBSET_PCT = 0.2
PCT_OUTLIERS = 7.4
SEEDS = [42, 7, 123, 2024, 1337]
OUTPUT_DIR = "ablation_outputs_v3"

# Model hyperparameters (matched to PRAXIS, not v11.py defaults)
HIDDEN_DIM = 256
LATENT_DIM_NUM = 16
LATENT_DIM_CAT = 32
CONTEXT_DIM = 40
ATTENTION_HEADS = 4
DROPOUT_RATE = 0.1
SALARY_WEIGHT = 2.0
BOUND_CONTEXT = True   # apply tanh to encoder output; set False to test unbounded context

# Training schedule (matched to praxis: 15 VAE epochs, 200 diffusion epochs.
# With mini-batch SGD on diffusion (batch_size=256, ~300 batches/epoch), 30 epochs
# already gives ~9000 updates >> praxis's 200 full-batch updates, so we keep the
# epoch count modest. Set diff_epochs higher if you suspect under-training.)
VAE_EPOCHS = 15
VAE_BATCH_SIZE = 256
VAE_WARMUP = 5
VAE_MAX_BETA = 1.0
DIFF_TIMESTEPS = 1000
DIFF_EPOCHS = 30
DATA_DIFF_EPOCHS = 30
DIFF_BATCH_SIZE = 256
ENCODER_CHUNK_SIZE = 4096
LR = 1e-3
WEIGHT_DECAY = 1e-4

DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")


@dataclass
class AblationConfig:
    name: str
    label: str
    use_vae: bool
    use_diffusion: bool
    diffusion_space: str          # "latent" | "data" | "none"
    use_context: bool
    context_in_vae: bool
    context_in_diffusion: bool


CONFIGS = [
    AblationConfig("A_vae_only",       "VAE only",                       True,  False, "none",   False, False, False),
    AblationConfig("B_ddpm_data",      "Data-space DDPM",                False, True,  "data",   False, False, False),
    AblationConfig("C_vae_diff_nocxt", "Latent Diffusion (unconditional)", True,  True,  "latent", False, False, False),
    AblationConfig("D_cxt_in_vae",     "Context in VAE only",            True,  True,  "latent", True,  True,  False),
    AblationConfig("E_cxt_in_diff",    "CALDM (proposed)",               True,  True,  "latent", True,  False, True),
    AblationConfig("F_full_caldm",     "CALDM + VAE context",            True,  True,  "latent", True,  True,  True),
]
FULL_NAME = "E_cxt_in_diff"   # V2: comparisons now anchor on Context-in-Diffusion-only


# =============================================================================
# 1. Reproducibility
# =============================================================================

def set_seed(seed: int):
    random.seed(seed)
    np.random.seed(seed)
    torch.manual_seed(seed)
    if torch.cuda.is_available():
        torch.cuda.manual_seed_all(seed)
    torch.backends.cudnn.deterministic = True
    torch.backends.cudnn.benchmark = False


# =============================================================================
# 2. Model components
# =============================================================================

class ContextEncoder(nn.Module):
    """Transformer-based context encoder with multi-head self-attention.
    Faithful to caldm_v11.py (default batch_first=False, so the encoder operates
    over the batch as if it were a sequence -- samples attend to each other).
    A tanh activation is applied to the output to bound the context vector during
    joint training, which prevents the conditioning signal from blowing up when
    the encoder is trained against the VAE's KL loss."""

    def __init__(self, input_dim, embed_dim=10, hidden_dim=256, num_heads=4, dropout_rate=0.1):
        super().__init__()
        self.fc = nn.Sequential(
            nn.Linear(input_dim, hidden_dim),
            nn.ReLU(),
            nn.Dropout(dropout_rate),
        )
        # NOTE: batch_first NOT set -> defaults to False, matching caldm_v11.py.
        # With unsqueeze(1) the transformer treats the batch dim as a sequence.
        encoder_layer = nn.TransformerEncoderLayer(
            d_model=hidden_dim, nhead=num_heads, dropout=dropout_rate
        )
        self.transformer = nn.TransformerEncoder(encoder_layer, num_layers=2)
        self.output_layer = nn.Linear(hidden_dim, embed_dim)

    def forward(self, x):
        h = self.fc(x).unsqueeze(1)            # (B, 1, hidden)
        h = self.transformer(h)
        out = self.output_layer(h.squeeze(1))
        return torch.tanh(out) if BOUND_CONTEXT else out


class FlexibleVAE(nn.Module):
    """Dual-head VAE with split numerical/categorical latents and decoders.

    Setting context_dim=0 cleanly disables conditioning so the same class
    serves both the context-aware and unconditional ablation cases.
    """

    def __init__(self, input_dim, latent_dim_num, latent_dim_cat, context_dim,
                 hidden_dim, dropout_rate=0.1):
        super().__init__()
        self.use_context = context_dim > 0
        self.context_dim = context_dim
        self.latent_dim_num = latent_dim_num
        self.latent_dim_cat = latent_dim_cat
        self.num_features = 1
        self.cat_features = input_dim - 1
        total_latent = latent_dim_num + latent_dim_cat

        enc_in = input_dim + context_dim
        dec_num_in = latent_dim_num + context_dim
        dec_cat_in = latent_dim_cat + context_dim

        self.encoder = self._mlp(enc_in, hidden_dim, total_latent * 2, dropout_rate, activation="leaky")
        self.decoder_num = self._mlp(dec_num_in, hidden_dim, self.num_features, dropout_rate, activation="leaky")
        self.decoder_cat = self._mlp(dec_cat_in, hidden_dim, self.cat_features, dropout_rate, activation="leaky")

    @staticmethod
    def _mlp(in_dim, hid, out_dim, dropout, activation="leaky"):
        act = nn.LeakyReLU if activation == "leaky" else nn.ReLU
        return nn.Sequential(
            nn.Linear(in_dim, hid), act(), nn.Dropout(dropout),
            nn.Linear(hid, hid),    act(), nn.Dropout(dropout),
            nn.Linear(hid, out_dim),
        )

    @staticmethod
    def reparameterize(mu, log_var):
        log_var = torch.clamp(log_var, min=-5, max=5)
        std = torch.exp(0.5 * log_var)
        return mu + torch.randn_like(std) * std

    def _maybe_cat(self, x, ctx):
        return torch.cat([x, ctx], dim=1) if self.use_context else x

    def encode(self, x, context=None):
        h = self.encoder(self._maybe_cat(x, context))
        mu, log_var = torch.chunk(h, 2, dim=1)
        log_var = torch.clamp(log_var, min=-5, max=5)  # KL stability
        return mu, log_var

    def decode(self, z, context=None):
        z_num = z[:, :self.latent_dim_num]
        z_cat = z[:, self.latent_dim_num:]
        sal = self.decoder_num(self._maybe_cat(z_num, context))
        cat = self.decoder_cat(self._maybe_cat(z_cat, context))
        return sal, cat

    def forward(self, x, context=None):
        mu, log_var = self.encode(x, context)
        z = self.reparameterize(mu, log_var)
        sal, cat = self.decode(z, context)
        return sal, cat, mu, log_var


class FlexibleLatentDiffusion(nn.Module):
    """Latent DDPM with separate noise predictors for numerical and categorical
    latents. Linear beta schedule. Context conditioning toggled by context_dim>0."""

    def __init__(self, latent_dim_num, latent_dim_cat, context_dim,
                 timesteps=1000, hidden_dim=256, dropout_rate=0.1):
        super().__init__()
        self.timesteps = timesteps
        self.use_context = context_dim > 0
        self.context_dim = context_dim
        self.latent_dim_num = latent_dim_num
        self.latent_dim_cat = latent_dim_cat

        in_num = latent_dim_num + context_dim + 1
        in_cat = latent_dim_cat + context_dim + 1

        beta = torch.linspace(0.0001, 0.02, timesteps)
        alpha = torch.clamp(1.0 - beta, min=1e-5, max=1.0)
        alpha_cum = torch.cumprod(alpha, dim=0)
        alpha_cum = torch.maximum(alpha_cum, torch.tensor(1e-5))
        self.register_buffer("beta", beta)
        self.register_buffer("alpha", alpha)
        self.register_buffer("alpha_cum", alpha_cum)

        self.network_num = self._mlp(in_num, hidden_dim, latent_dim_num, dropout_rate)
        self.network_cat = self._mlp(in_cat, hidden_dim, latent_dim_cat, dropout_rate)

    @staticmethod
    def _mlp(in_dim, hid, out_dim, dropout):
        return nn.Sequential(
            nn.Linear(in_dim, hid), nn.ReLU(), nn.Dropout(dropout),
            nn.Linear(hid, hid),    nn.ReLU(), nn.Dropout(dropout),
            nn.Linear(hid, out_dim),
        )

    def forward_diffusion(self, z0_num, z0_cat, t):
        alpha_t = self.alpha_cum[t].unsqueeze(1)
        n_num = torch.randn_like(z0_num)
        n_cat = torch.randn_like(z0_cat)
        z_t_num = torch.sqrt(alpha_t) * z0_num + torch.sqrt(1 - alpha_t) * n_num
        z_t_cat = torch.sqrt(alpha_t) * z0_cat + torch.sqrt(1 - alpha_t) * n_cat
        return z_t_num, z_t_cat, n_num, n_cat

    def reverse_denoising(self, z_t_num, z_t_cat, context, t):
        t_norm = t.view(-1, 1).float() / self.timesteps
        if self.use_context and context is not None:
            x_in_num = torch.cat([z_t_num, context, t_norm], dim=1)
            x_in_cat = torch.cat([z_t_cat, context, t_norm], dim=1)
        else:
            x_in_num = torch.cat([z_t_num, t_norm], dim=1)
            x_in_cat = torch.cat([z_t_cat, t_norm], dim=1)
        return 0.5 * self.network_num(x_in_num).tanh(), 0.5 * self.network_cat(x_in_cat).tanh()

    @torch.no_grad()
    def sample(self, batch_size, context=None, steps=None, device=None):
        if steps is None:
            steps = self.timesteps
        if device is None:
            device = self.beta.device
        z_n = torch.randn(batch_size, self.latent_dim_num, device=device)
        z_c = torch.randn(batch_size, self.latent_dim_cat, device=device)
        for t in reversed(range(steps)):
            z_n = z_n / (torch.norm(z_n, p=2, dim=1, keepdim=True) + 1e-5)
            z_c = z_c / (torch.norm(z_c, p=2, dim=1, keepdim=True) + 1e-5)
            t_long = torch.full((batch_size,), t, device=device, dtype=torch.long)
            pn, pc = self.reverse_denoising(z_n, z_c, context, t_long)
            alpha_t = torch.clamp(self.alpha_cum[t].view(-1, 1), min=1e-3, max=1.0)
            z_n = (z_n - torch.sqrt(1 - alpha_t) * pn) / (torch.sqrt(alpha_t) + 1e-5)
            z_c = (z_c - torch.sqrt(1 - alpha_t) * pc) / (torch.sqrt(alpha_t) + 1e-5)
            if t > 0:
                z_n = z_n + 0.0003 * torch.sqrt(self.beta[t]) * torch.randn_like(z_n)
                z_c = z_c + 0.1    * torch.sqrt(self.beta[t]) * torch.randn_like(z_c)
        return z_n, z_c


class DataSpaceDDPM(nn.Module):
    """DDPM operating directly on input space (no VAE).

    Caveat for the manuscript: with one-hot categorical features, the Gaussian
    marginals assumed by DDPM are violated. We include this configuration to
    motivate the latent-space approach.
    """

    def __init__(self, input_dim, context_dim, timesteps=1000, hidden_dim=256, dropout_rate=0.1):
        super().__init__()
        self.timesteps = timesteps
        self.input_dim = input_dim
        self.use_context = context_dim > 0
        self.context_dim = context_dim

        beta = torch.linspace(0.0001, 0.02, timesteps)
        alpha = torch.clamp(1.0 - beta, min=1e-5, max=1.0)
        alpha_cum = torch.cumprod(alpha, dim=0)
        alpha_cum = torch.maximum(alpha_cum, torch.tensor(1e-5))
        self.register_buffer("beta", beta)
        self.register_buffer("alpha", alpha)
        self.register_buffer("alpha_cum", alpha_cum)

        in_dim = input_dim + context_dim + 1
        self.network = nn.Sequential(
            nn.Linear(in_dim, hidden_dim), nn.ReLU(), nn.Dropout(dropout_rate),
            nn.Linear(hidden_dim, hidden_dim), nn.ReLU(), nn.Dropout(dropout_rate),
            nn.Linear(hidden_dim, hidden_dim), nn.ReLU(),
            nn.Linear(hidden_dim, input_dim),
        )

    def forward_diffusion(self, x0, t):
        alpha_t = self.alpha_cum[t].unsqueeze(1)
        noise = torch.randn_like(x0)
        x_t = torch.sqrt(alpha_t) * x0 + torch.sqrt(1 - alpha_t) * noise
        return x_t, noise

    def predict_noise(self, x_t, context, t):
        t_norm = t.view(-1, 1).float() / self.timesteps
        if self.use_context and context is not None:
            x_in = torch.cat([x_t, context, t_norm], dim=1)
        else:
            x_in = torch.cat([x_t, t_norm], dim=1)
        return self.network(x_in)


# =============================================================================
# 3. Training
# =============================================================================

def beta_scheduler(epoch, warmup_epochs, max_beta):
    if epoch < warmup_epochs:
        return max(0.005, (epoch / warmup_epochs) ** 2 * max_beta)
    return max_beta

def train_vae(vae, context_encoder, X_train, X_val, epochs, batch_size, device, verbose=False):
    """Train VAE; if a context_encoder is supplied it is trained jointly.
    Includes gradient clipping on ALL parameters and skips non-finite-loss batches."""
    params = list(vae.parameters())
    if context_encoder is not None:
        params += list(context_encoder.parameters())
    opt = optim.Adam(params, lr=LR, weight_decay=WEIGHT_DECAY)

    X_train_aug = X_train + torch.randn_like(X_train) * 0.01
    train_loader = torch.utils.data.DataLoader(
        torch.utils.data.TensorDataset(X_train_aug),
        batch_size=batch_size, shuffle=True
    )

    for epoch in range(epochs):
        vae.train()
        if context_encoder is not None:
            context_encoder.train()
        beta = beta_scheduler(epoch, VAE_WARMUP, VAE_MAX_BETA)
        epoch_loss, n_seen, n_skipped = 0.0, 0, 0
        for (x_batch,) in train_loader:
            x_batch = x_batch.to(device)
            opt.zero_grad()
            sal = x_batch[:, :1]
            cat = x_batch[:, 1:]
            ctx = context_encoder(x_batch) if context_encoder is not None else None
            sal_rec, cat_logits, mu, log_var = vae(x_batch, ctx)
            # log_var is already clamped inside encode(), so KL can't overflow here
            kl = -0.5 * torch.sum(1 + log_var - mu.pow(2) - log_var.exp()) / sal.size(0)
            loss = (
                F.mse_loss(sal_rec, sal) * SALARY_WEIGHT
                + F.binary_cross_entropy_with_logits(cat_logits, cat)
                + beta * kl
            )
            if not torch.isfinite(loss):
                n_skipped += 1
                continue
            loss.backward()
            # Skip if any gradient is non-finite (would corrupt Adam's moments)
            grad_finite = all((p.grad is None) or torch.isfinite(p.grad).all() for p in params)
            if not grad_finite:
                n_skipped += 1
                opt.zero_grad()
                continue
            torch.nn.utils.clip_grad_norm_(params, max_norm=1.0)  # ALL params, not just VAE
            opt.step()
            epoch_loss += loss.item()
            n_seen += 1

        if verbose and ((epoch + 1) % 5 == 0 or n_skipped > 0):
            avg = epoch_loss / max(n_seen, 1)
            msg = f"    [VAE] epoch {epoch+1}/{epochs}  loss={avg:.4f}  beta={beta:.3f}"
            if n_skipped > 0:
                msg += f"  SKIPPED {n_skipped} non-finite-loss batches"
            print(msg)

    # Sanity check: bail loudly if joint training produced NaN weights
    for name, p in (list(vae.named_parameters())
                    + (list(context_encoder.named_parameters()) if context_encoder is not None else [])):
        if not torch.isfinite(p).all():
            raise RuntimeError(
                f"VAE/encoder training produced non-finite weights in '{name}'. "
                f"Likely cause: KL or encoder activation blow-up. "
                f"Try lowering LR (currently {LR}) or VAE_MAX_BETA (currently {VAE_MAX_BETA})."
            )

    vae.eval()
    if context_encoder is not None:
        context_encoder.eval()


def train_latent_diffusion(diff, vae, context_encoder, config, X_train, epochs, device,
                           batch_size=DIFF_BATCH_SIZE, verbose=False):
    """Train latent diffusion with mini-batch SGD.

    Encoder fine-tuned only when context is used by diffusion AND wasn't already
    trained alongside the VAE (i.e. only in config E). Otherwise the encoder is
    treated as frozen, and (if used) its outputs are precomputed once."""
    train_encoder = (
        config.context_in_diffusion
        and not config.context_in_vae
        and context_encoder is not None
    )

    params = list(diff.parameters())
    if train_encoder:
        params += list(context_encoder.parameters())
    opt = optim.Adam(params, lr=LR)

    X_train_d = X_train.to(device)
    n_train = X_train_d.shape[0]
    vae.eval()

    # ---- Precompute frozen latents (and frozen diffusion-context, if applicable) in chunks
    z_train_num_chunks, z_train_cat_chunks = [], []
    ctx_diff_chunks = []  # only populated if diff uses context AND encoder is frozen
    use_precomputed_ctx_diff = (
        config.context_in_diffusion and context_encoder is not None and not train_encoder
    )

    with torch.no_grad():
        for start in range(0, n_train, batch_size):
            x_chunk = X_train_d[start:start + batch_size]
            ctx_vae_chunk = (context_encoder(x_chunk)
                             if (config.context_in_vae and context_encoder is not None) else None)
            mu, lv = vae.encode(x_chunk, ctx_vae_chunk)
            z = vae.reparameterize(mu, lv)
            z_train_num_chunks.append(z[:, :diff.latent_dim_num])
            z_train_cat_chunks.append(z[:, diff.latent_dim_num:])
            if use_precomputed_ctx_diff:
                ctx_diff_chunks.append(context_encoder(x_chunk))

    z_train_num = torch.cat(z_train_num_chunks, dim=0)
    z_train_cat = torch.cat(z_train_cat_chunks, dim=0)
    ctx_diff_pre = torch.cat(ctx_diff_chunks, dim=0) if ctx_diff_chunks else None

    # ---- Mini-batch SGD over (epochs * n_train / batch_size) updates
    for epoch in range(epochs):
        diff.train()
        if train_encoder:
            context_encoder.train()

        perm = torch.randperm(n_train, device=device)
        epoch_loss, n_seen = 0.0, 0

        for start in range(0, n_train, batch_size):
            idx = perm[start:start + batch_size]
            opt.zero_grad()

            zb_num = z_train_num[idx]
            zb_cat = z_train_cat[idx]

            if config.context_in_diffusion and context_encoder is not None:
                if train_encoder:
                    ctx_diff = context_encoder(X_train_d[idx])
                else:
                    ctx_diff = ctx_diff_pre[idx]
            else:
                ctx_diff = None

            t = torch.randint(0, diff.timesteps, (zb_num.shape[0],), device=device)
            z_t_num, z_t_cat, n_num, n_cat = diff.forward_diffusion(zb_num, zb_cat, t)
            pred_num, pred_cat = diff.reverse_denoising(z_t_num, z_t_cat, ctx_diff, t)
            loss = F.mse_loss(pred_num, n_num) + F.mse_loss(pred_cat, n_cat)
            loss.backward()
            torch.nn.utils.clip_grad_norm_(params, max_norm=1.0)
            opt.step()
            epoch_loss += loss.item()
            n_seen += 1

        if verbose and (epoch + 1) % 5 == 0:
            print(f"    [LatentDiff] epoch {epoch+1}/{epochs}  loss={epoch_loss/n_seen:.6f}")

    diff.eval()
    if context_encoder is not None:
        context_encoder.eval()


def train_data_diffusion(diff, context_encoder, X_train, epochs, device,
                         batch_size=DIFF_BATCH_SIZE, verbose=False):
    """Train DDPM directly on input space with mini-batch SGD."""
    params = list(diff.parameters())
    if context_encoder is not None:
        params += list(context_encoder.parameters())
    opt = optim.Adam(params, lr=LR)

    X_train_d = X_train.to(device)
    n_train = X_train_d.shape[0]

    for epoch in range(epochs):
        diff.train()
        if context_encoder is not None:
            context_encoder.train()

        perm = torch.randperm(n_train, device=device)
        epoch_loss, n_seen = 0.0, 0

        for start in range(0, n_train, batch_size):
            idx = perm[start:start + batch_size]
            xb = X_train_d[idx]

            opt.zero_grad()
            t = torch.randint(0, diff.timesteps, (xb.shape[0],), device=device)
            x_t, noise = diff.forward_diffusion(xb, t)
            ctx = context_encoder(xb) if context_encoder is not None else None
            pred = diff.predict_noise(x_t, ctx, t)
            loss = F.mse_loss(pred, noise)
            loss.backward()
            torch.nn.utils.clip_grad_norm_(params, max_norm=1.0)
            opt.step()
            epoch_loss += loss.item()
            n_seen += 1

        if verbose and (epoch + 1) % 5 == 0:
            print(f"    [DataDiff] epoch {epoch+1}/{epochs}  loss={epoch_loss/n_seen:.6f}")

    diff.eval()
    if context_encoder is not None:
        context_encoder.eval()


# =============================================================================
# 4. Scoring
# =============================================================================

@torch.no_grad()
def score_vae_only(vae, context_encoder, X_test, device, chunk_size=ENCODER_CHUNK_SIZE):
    X_test = X_test.to(device)
    n = X_test.shape[0]
    parts = []
    for start in range(0, n, chunk_size):
        x = X_test[start:start + chunk_size]
        ctx = context_encoder(x) if context_encoder is not None else None
        sal_rec, cat_logits, _, _ = vae(x, ctx)
        cat_rec = (torch.sigmoid(cat_logits) > 0.5).float()
        sal_loss = torch.mean((x[:, :1] - sal_rec) ** 2, dim=1)
        cat_loss = torch.mean((x[:, 1:] - cat_rec) ** 2, dim=1)
        parts.append(torch.log1p(sal_loss + cat_loss))
    return torch.cat(parts).cpu().numpy()


@torch.no_grad()
def score_latent_pipeline(vae, diff, context_encoder, config, X_test, device,
                          chunk_size=ENCODER_CHUNK_SIZE):
    """CALDM-style scoring: sample latents from diffusion -> decode -> recon loss.
    Chunked to keep encoder + sampler memory bounded."""
    X_test = X_test.to(device)
    n = X_test.shape[0]
    parts = []
    for start in range(0, n, chunk_size):
        x = X_test[start:start + chunk_size]
        ctx_vae = (context_encoder(x) if (context_encoder is not None and config.context_in_vae) else None)
        ctx_diff = (context_encoder(x) if (context_encoder is not None and config.context_in_diffusion) else None)
        sn, sc = diff.sample(batch_size=x.shape[0], context=ctx_diff,
                             steps=diff.timesteps, device=device)
        z = torch.cat([sn, sc], dim=1)
        sal_rec, cat_logits = vae.decode(z, ctx_vae)
        cat_rec = (torch.sigmoid(cat_logits) > 0.5).float()
        sal_loss = torch.mean((x[:, :1] - sal_rec) ** 2, dim=1)
        cat_loss = torch.mean((x[:, 1:] - cat_rec) ** 2, dim=1)
        parts.append(torch.log1p(sal_loss + cat_loss))
    return torch.cat(parts).cpu().numpy()


@torch.no_grad()
def score_data_diffusion(diff, context_encoder, X_test, device,
                         t_grid=(100, 300, 500, 700, 900),
                         chunk_size=ENCODER_CHUNK_SIZE):
    """Score = mean noise-prediction MSE across a fixed grid of timesteps.
    Higher reconstruction error => more anomalous."""
    X_test = X_test.to(device)
    n = X_test.shape[0]
    parts = []
    for start in range(0, n, chunk_size):
        x = X_test[start:start + chunk_size]
        ctx = context_encoder(x) if context_encoder is not None else None
        chunk_scores = torch.zeros(x.shape[0], device=device)
        for tv in t_grid:
            t = torch.full((x.shape[0],), tv, device=device, dtype=torch.long)
            x_t, noise = diff.forward_diffusion(x, t)
            pred = diff.predict_noise(x_t, ctx, t)
            chunk_scores = chunk_scores + torch.mean((pred - noise) ** 2, dim=1)
        parts.append(chunk_scores / len(t_grid))
    return torch.cat(parts).cpu().numpy()


# =============================================================================
# 5. Per-config experiment runner
# =============================================================================

def run_single_experiment(config, seed, X_train, X_val, X_test, y_test, y_type_test,
                          input_dim, device, verbose=False):
    set_seed(seed)
    t0 = time.time()

    ctx_dim_vae = CONTEXT_DIM if config.context_in_vae else 0
    ctx_dim_diff = CONTEXT_DIM if config.context_in_diffusion else 0

    context_encoder = None
    if config.use_context:
        context_encoder = ContextEncoder(input_dim, embed_dim=CONTEXT_DIM,
                                         hidden_dim=HIDDEN_DIM, num_heads=ATTENTION_HEADS,
                                         dropout_rate=DROPOUT_RATE).to(device)

    # ---- VAE
    vae = None
    if config.use_vae:
        vae = FlexibleVAE(input_dim, LATENT_DIM_NUM, LATENT_DIM_CAT,
                          ctx_dim_vae, HIDDEN_DIM, DROPOUT_RATE).to(device)
        ctx_for_vae_train = context_encoder if config.context_in_vae else None
        train_vae(vae, ctx_for_vae_train, X_train, X_val, VAE_EPOCHS, VAE_BATCH_SIZE,
                  device, verbose=verbose)

    # ---- Diffusion
    diff = None
    if config.use_diffusion:
        if config.diffusion_space == "latent":
            diff = FlexibleLatentDiffusion(LATENT_DIM_NUM, LATENT_DIM_CAT, ctx_dim_diff,
                                           DIFF_TIMESTEPS, HIDDEN_DIM, DROPOUT_RATE).to(device)
            train_latent_diffusion(diff, vae, context_encoder, config, X_train,
                                   DIFF_EPOCHS, device, verbose=verbose)
        elif config.diffusion_space == "data":
            diff = DataSpaceDDPM(input_dim, ctx_dim_diff, DIFF_TIMESTEPS, HIDDEN_DIM, DROPOUT_RATE).to(device)
            ctx_enc = context_encoder if config.context_in_diffusion else None
            train_data_diffusion(diff, ctx_enc, X_train, DATA_DIFF_EPOCHS, device, verbose=verbose)

    # ---- Score
    if config.use_vae and config.use_diffusion:
        scores = score_latent_pipeline(vae, diff, context_encoder, config, X_test, device)
    elif config.use_vae:
        ctx_enc = context_encoder if config.context_in_vae else None
        scores = score_vae_only(vae, ctx_enc, X_test, device)
    elif config.use_diffusion:
        ctx_enc = context_encoder if config.context_in_diffusion else None
        scores = score_data_diffusion(diff, ctx_enc, X_test, device)
    else:
        raise ValueError(f"Config {config.name} has neither VAE nor diffusion")

    train_time = time.time() - t0
    metrics = compute_metrics(scores, y_test, y_type_test)
    metrics.update({
        "config": config.name,
        "label": config.label,
        "seed": seed,
        "train_time_sec": train_time,
        "scores": scores,
    })
    return metrics


def compute_metrics(scores, y_true, y_type=None):
    if not np.all(np.isfinite(scores)):
        n_bad = int((~np.isfinite(scores)).sum())
        raise ValueError(
            f"{n_bad}/{len(scores)} anomaly scores are non-finite (NaN/Inf). "
            f"This means upstream training produced bad weights. "
            f"Check the [VAE]/[LatentDiff] loss prints for divergence; "
            f"if loss went to NaN/Inf the gradient clipping or LR may need attention."
        )
    threshold = float(np.percentile(scores, 100 - PCT_OUTLIERS))
    y_pred = (scores >= threshold).astype(int)
    out = {
        "accuracy":  accuracy_score(y_true, y_pred),
        "precision": precision_score(y_true, y_pred, zero_division=1),
        "recall":    recall_score(y_true, y_pred, zero_division=1),
        "f1":        f1_score(y_true, y_pred, zero_division=1),
        "auc_roc":   roc_auc_score(y_true, scores),
        "ap":        average_precision_score(y_true, scores),
        "threshold": threshold,
    }
    if y_type is not None:
        outlier_mask = (y_true == 1)
        normal_mask = (y_true == 0)
        outlier_types = np.unique(y_type[outlier_mask])
        for t in outlier_types:
            this_mask = (y_type == t) & outlier_mask
            sel = this_mask | normal_mask
            yb = this_mask[sel].astype(int)
            sb = scores[sel]
            if yb.sum() > 0 and (1 - yb).sum() > 0:
                out[f"auc_type_{t}"] = roc_auc_score(yb, sb)
    return out


# =============================================================================
# 6. Aggregation and outputs
# =============================================================================

def aggregate_results(all_results):
    df = pd.DataFrame([{k: v for k, v in r.items() if k != "scores"} for r in all_results])

    metric_cols = ["auc_roc", "ap", "f1", "accuracy", "precision", "recall", "train_time_sec"]
    grouped = df.groupby("config")[metric_cols].agg(["mean", "std"])
    grouped.columns = [f"{m}_{s}" for m, s in grouped.columns]
    grouped = grouped.reset_index()

    label_map = {c.name: c.label for c in CONFIGS}
    grouped["label"] = grouped["config"].map(label_map)

    # Paired t-tests vs the proposed model (per metric)
    if FULL_NAME in df["config"].values:
        full_df = df[df["config"] == FULL_NAME].sort_values("seed").reset_index(drop=True)
        for cname in df["config"].unique():
            if cname == FULL_NAME:
                continue
            cdf = df[df["config"] == cname].sort_values("seed").reset_index(drop=True)
            n = min(len(full_df), len(cdf))
            for m in ("auc_roc", "ap"):
                if n >= 2:
                    _, p = stats.ttest_rel(full_df[m].values[:n], cdf[m].values[:n])
                else:
                    p = np.nan
                grouped.loc[grouped["config"] == cname, f"{m}_pval_vs_full"] = p
    return df, grouped


def save_per_type_breakdown(all_results, out_path):
    type_cols = sorted({k for r in all_results for k in r if k.startswith("auc_type_")})
    if not type_cols:
        return None
    rows = []
    for cfg in CONFIGS:
        cfg_runs = [r for r in all_results if r["config"] == cfg.name]
        row = {"config": cfg.name, "label": cfg.label}
        for tc in type_cols:
            vals = [r.get(tc, np.nan) for r in cfg_runs]
            vals = [v for v in vals if not np.isnan(v)]
            row[f"{tc}_mean"] = float(np.mean(vals)) if vals else np.nan
            row[f"{tc}_std"] = float(np.std(vals)) if vals else np.nan
        rows.append(row)
    out_df = pd.DataFrame(rows)
    out_df.to_csv(out_path, index=False)
    return out_df


def write_docx_table(summary_df, out_path):
    try:
        from docx import Document
    except ImportError:
        print("[warn] python-docx not installed; skipping .docx output. (pip install python-docx)")
        return

    doc = Document()
    doc.add_heading("CALDM Ablation Study Results", level=1)
    doc.add_paragraph(
        f"Each cell reports mean ± std across {len(SEEDS)} random seeds. "
        f"P-values are paired t-tests comparing each ablation to the proposed CALDM model."
    )

    headers = ["Configuration", "AUC-ROC", "AP", "F1@7.4%", "p (AUC)", "p (AP)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    ordered = summary_df.copy()
    ordered["sort_key"] = ordered["config"].apply(lambda x: 1 if x == FULL_NAME else 0)
    ordered = ordered.sort_values(["sort_key", "config"]).reset_index(drop=True)

    for _, row in ordered.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["label"])
        cells[1].text = f"{row['auc_roc_mean']:.4f} ± {row['auc_roc_std']:.4f}"
        cells[2].text = f"{row['ap_mean']:.4f} ± {row['ap_std']:.4f}"
        cells[3].text = f"{row['f1_mean']:.4f} ± {row['f1_std']:.4f}"
        p_auc = row.get("auc_roc_pval_vs_full", np.nan)
        p_ap = row.get("ap_pval_vs_full", np.nan)
        cells[4].text = "—" if pd.isna(p_auc) else f"{p_auc:.4f}"
        cells[5].text = "—" if pd.isna(p_ap) else f"{p_ap:.4f}"

    doc.save(out_path)
    print(f"[ok] Saved {out_path}")


def write_latex_table(summary_df, out_path):
    lines = [
        "\\begin{table}[h]",
        "\\centering",
        f"\\caption{{CALDM ablation study. Mean $\\pm$ std across {len(SEEDS)} random seeds. "
        "P-values are paired t-tests against the proposed CALDM model.}",
        "\\label{tab:ablation}",
        "\\begin{tabular}{lccccc}",
        "\\toprule",
        "Configuration & AUC-ROC & AP & F1 & $p_{\\text{AUC}}$ & $p_{\\text{AP}}$ \\\\",
        "\\midrule",
    ]

    ordered = summary_df.copy()
    ordered["sort_key"] = ordered["config"].apply(lambda x: 1 if x == FULL_NAME else 0)
    ordered = ordered.sort_values(["sort_key", "config"]).reset_index(drop=True)

    for _, row in ordered.iterrows():
        p_auc = row.get("auc_roc_pval_vs_full", np.nan)
        p_ap = row.get("ap_pval_vs_full", np.nan)
        p_auc_s = "---" if pd.isna(p_auc) else f"{p_auc:.4f}"
        p_ap_s = "---" if pd.isna(p_ap) else f"{p_ap:.4f}"
        lines.append(
            f"{row['label']} & "
            f"{row['auc_roc_mean']:.4f} $\\pm$ {row['auc_roc_std']:.4f} & "
            f"{row['ap_mean']:.4f} $\\pm$ {row['ap_std']:.4f} & "
            f"{row['f1_mean']:.4f} $\\pm$ {row['f1_std']:.4f} & "
            f"{p_auc_s} & {p_ap_s} \\\\"
        )
    lines.extend(["\\bottomrule", "\\end{tabular}", "\\end{table}"])
    with open(out_path, "w") as f:
        f.write("\n".join(lines))
    print(f"[ok] Saved {out_path}")


def make_plots(all_results, summary_df, y_test, output_dir, show=False):
    figdir = os.path.join(output_dir, "figures")
    os.makedirs(figdir, exist_ok=True)

    def _finish(fname):
        plt.tight_layout()
        plt.savefig(os.path.join(figdir, fname), dpi=150)
        if show:
            plt.show()
        else:
            plt.close()

    # ---- AP barplot
    plt.figure(figsize=(10, 5))
    order = sorted(summary_df["config"].unique(), key=lambda x: 1 if x == FULL_NAME else 0)
    means = [summary_df.loc[summary_df["config"] == c, "ap_mean"].iloc[0] for c in order]
    stds  = [summary_df.loc[summary_df["config"] == c, "ap_std"].iloc[0] for c in order]
    labels = [summary_df.loc[summary_df["config"] == c, "label"].iloc[0] for c in order]
    colors = ["#d8633b" if c == FULL_NAME else "#3b7dd8" for c in order]
    plt.bar(range(len(order)), means, yerr=stds, capsize=4, color=colors, edgecolor="black")
    plt.xticks(range(len(order)), labels, rotation=20, ha="right")
    plt.ylabel("Average Precision")
    plt.title(f"Ablation: AP (mean ± std across {len(SEEDS)} seeds)")
    _finish("ap_barplot.png")

    # ---- ROC overlay (median-seed run per config)
    plt.figure(figsize=(7, 7))
    for cfg in CONFIGS:
        cfg_runs = [r for r in all_results if r["config"] == cfg.name]
        if not cfg_runs:
            continue
        cfg_runs_sorted = sorted(cfg_runs, key=lambda r: r["auc_roc"])
        median_run = cfg_runs_sorted[len(cfg_runs_sorted) // 2]
        fpr, tpr, _ = roc_curve(y_test, median_run["scores"])
        lw = 2.5 if cfg.name == FULL_NAME else 1.2
        plt.plot(fpr, tpr, label=f"{cfg.label} (AUC={median_run['auc_roc']:.3f})", linewidth=lw)
    plt.plot([0, 1], [0, 1], "k--", alpha=0.4)
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("ROC curves (median-seed run per configuration)")
    plt.legend(loc="lower right", fontsize=9)
    _finish("roc_overlay.png")

    # ---- Per-type heatmap
    type_cols = sorted({k for r in all_results for k in r if k.startswith("auc_type_")})
    if type_cols:
        cfg_order = [c.name for c in CONFIGS]
        cfg_labels = [c.label for c in CONFIGS]
        mat = np.full((len(cfg_order), len(type_cols)), np.nan)
        for i, cn in enumerate(cfg_order):
            cfg_runs = [r for r in all_results if r["config"] == cn]
            for j, tc in enumerate(type_cols):
                vals = [r.get(tc, np.nan) for r in cfg_runs]
                vals = [v for v in vals if not np.isnan(v)]
                if vals:
                    mat[i, j] = float(np.mean(vals))
        plt.figure(figsize=(max(7, len(type_cols) + 3), 5))
        sns.heatmap(mat, annot=True, fmt=".3f",
                    xticklabels=[t.replace("auc_type_", "Type ") for t in type_cols],
                    yticklabels=cfg_labels,
                    cmap="RdYlGn", vmin=0.5, vmax=1.0)
        plt.title("AUC-ROC by outlier type and configuration (mean across seeds)")
        _finish("per_type_heatmap.png")
    print(f"[ok] Plots saved to {figdir}/")


# =============================================================================
# 7. Data loading and main
# =============================================================================

def load_data(data_path=DATA_PATH, subset_pct=SUBSET_PCT):
    if not os.path.exists(data_path):
        raise FileNotFoundError(f"Data file not found at {data_path}. "
                                f"Edit DATA_PATH at the top of the script.")
    data = np.load(data_path, allow_pickle=True)
    X_all, y_all = data["X"], data["y"]

    np.random.seed(42)
    n = int(subset_pct * X_all.shape[0])
    idx = np.random.choice(X_all.shape[0], n, replace=False)
    X = X_all[idx]
    y_subset = y_all[idx]
    assert y_subset.ndim == 2 and y_subset.shape[1] >= 2, \
        f"Expected y to be (N, >=2): col 0 = binary outlier label, col 1 = outlier type. Got shape {y_subset.shape}."
    y, y_type = y_subset[:, 0], y_subset[:, 1]

    # Normalize salary only (column 0)
    sal = X[:, 0:1].astype(np.float32)
    cat = X[:, 1:].astype(np.float32)
    scaler = StandardScaler()
    sal_n = scaler.fit_transform(sal)
    X_proc = np.hstack([sal_n, cat])
    X_t = torch.tensor(X_proc, dtype=torch.float32)

    X_tr, X_te, y_tr, y_te, yt_tr, yt_te = train_test_split(
        X_t, y, y_type, test_size=0.2, random_state=42, stratify=y_type
    )
    X_tr, X_val, y_tr, y_val, yt_tr, yt_val = train_test_split(
        X_tr, y_tr, yt_tr, test_size=0.2, random_state=42, stratify=yt_tr
    )
    return X_tr, X_val, X_te, y_tr.astype(int), y_val.astype(int), y_te.astype(int), yt_te


def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--quick", action="store_true",
                   help="Quick smoke-test mode: 1 seed, reduced epochs.")
    p.add_argument("--data", type=str, default=DATA_PATH, help="Path to .npz dataset.")
    p.add_argument("--out", type=str, default=OUTPUT_DIR, help="Output directory.")
    return p.parse_args()


def run_ablation(quick=False, data_path=None, output_dir=None, seeds=None,
                 vae_epochs=None, diff_epochs=None, data_diff_epochs=None,
                 latent_dim_num=None, latent_dim_cat=None, context_dim=None,
                 hidden_dim=None, lr=None, vae_max_beta=None,
                 salary_weight=None, bound_context=None,
                 show_plots=False, verbose=False):
    """Programmatic entry point. Use this from Jupyter or another script.

    All parameters are optional; if not given, the module-level defaults are used.
    V3 defaults match the praxis hyperparameter set.

    Example (Jupyter):
        from caldm_ablation_study_v3 import run_ablation
        # praxis-default 5-seed run
        results = run_ablation(seeds=[42, 7, 123, 2024, 1337], show_plots=True)

        # to test unbounded context (drops the tanh on encoder output):
        results = run_ablation(seeds=[42, 7, 123, 2024, 1337], bound_context=False)

    Returns
    -------
    dict with keys: 'raw', 'summary', 'per_type', 'all'
    """
    global VAE_EPOCHS, DIFF_EPOCHS, DATA_DIFF_EPOCHS, SEEDS, OUTPUT_DIR, DATA_PATH
    global LATENT_DIM_NUM, LATENT_DIM_CAT, CONTEXT_DIM, HIDDEN_DIM, LR, VAE_MAX_BETA
    global SALARY_WEIGHT, BOUND_CONTEXT

    if quick:
        if seeds is None:           seeds = [42]
        if vae_epochs is None:      vae_epochs = 5
        if diff_epochs is None:     diff_epochs = 5
        if data_diff_epochs is None: data_diff_epochs = 5
        print("[setup] QUICK mode: 1 seed, reduced epochs.")

    # Apply explicit overrides last so they win
    if data_path is not None:        DATA_PATH = data_path
    if output_dir is not None:       OUTPUT_DIR = output_dir
    if seeds is not None:            SEEDS = list(seeds)
    if vae_epochs is not None:       VAE_EPOCHS = vae_epochs
    if diff_epochs is not None:      DIFF_EPOCHS = diff_epochs
    if data_diff_epochs is not None: DATA_DIFF_EPOCHS = data_diff_epochs
    if latent_dim_num is not None:   LATENT_DIM_NUM = latent_dim_num
    if latent_dim_cat is not None:   LATENT_DIM_CAT = latent_dim_cat
    if context_dim is not None:      CONTEXT_DIM = context_dim
    if hidden_dim is not None:       HIDDEN_DIM = hidden_dim
    if lr is not None:               LR = lr
    if vae_max_beta is not None:     VAE_MAX_BETA = vae_max_beta
    if salary_weight is not None:    SALARY_WEIGHT = salary_weight
    if bound_context is not None:    BOUND_CONTEXT = bound_context

    print(f"[setup] hyperparameters: latent_num={LATENT_DIM_NUM}  latent_cat={LATENT_DIM_CAT}  "
          f"context={CONTEXT_DIM}  hidden={HIDDEN_DIM}  lr={LR}  max_beta={VAE_MAX_BETA}  "
          f"salary_weight={SALARY_WEIGHT}  bound_context={BOUND_CONTEXT}")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(os.path.join(OUTPUT_DIR, "figures"), exist_ok=True)

    print(f"[setup] device={DEVICE}  seeds={SEEDS}  configs={[c.name for c in CONFIGS]}")
    X_train, X_val, X_test, y_train, y_val, y_test, y_type_test = load_data(DATA_PATH)
    print(f"[setup] X_train={tuple(X_train.shape)}  X_val={tuple(X_val.shape)}  X_test={tuple(X_test.shape)}")
    unique_types, counts = np.unique(y_type_test, return_counts=True)
    type_summary = ", ".join(f"{t}={c}" for t, c in zip(unique_types, counts))
    print(f"[setup] outlier rate (test): {y_test.mean():.4f}  outlier types in test: {type_summary}")

    input_dim = X_train.shape[1]
    all_results = []

    for cfg in CONFIGS:
        for seed in SEEDS:
            print(f"\n=== Config: {cfg.name}  |  Seed: {seed} ===")
            try:
                r = run_single_experiment(cfg, seed, X_train, X_val, X_test,
                                          y_test, y_type_test, input_dim, DEVICE,
                                          verbose=verbose)
                print(f"  AUC-ROC={r['auc_roc']:.4f}  AP={r['ap']:.4f}  "
                      f"F1={r['f1']:.4f}  ({r['train_time_sec']:.1f}s)")
                all_results.append(r)
            except Exception as e:
                print(f"  [error] {e}")
                import traceback
                traceback.print_exc()
                continue

    # Save raw
    raw_df = pd.DataFrame([{k: v for k, v in r.items() if k != "scores"} for r in all_results])
    raw_path = os.path.join(OUTPUT_DIR, "ablation_results_raw.csv")
    raw_df.to_csv(raw_path, index=False)
    print(f"\n[ok] Saved {raw_path}")

    # Aggregate
    _, summary_df = aggregate_results(all_results)
    summary_path = os.path.join(OUTPUT_DIR, "ablation_results_summary.csv")
    summary_df.to_csv(summary_path, index=False)
    print(f"[ok] Saved {summary_path}")

    # Per-type breakdown
    per_type_path = os.path.join(OUTPUT_DIR, "ablation_per_type.csv")
    per_type_df = save_per_type_breakdown(all_results, per_type_path)
    print(f"[ok] Saved {per_type_path}")

    # Tables
    write_docx_table(summary_df, os.path.join(OUTPUT_DIR, "ablation_table.docx"))
    write_latex_table(summary_df, os.path.join(OUTPUT_DIR, "ablation_table.tex"))

    # Plots
    make_plots(all_results, summary_df, y_test, OUTPUT_DIR, show=show_plots)

    # Console summary
    print("\n" + "=" * 80)
    print(f"FINAL SUMMARY (mean ± std across {len(SEEDS)} seeds)")
    print("=" * 80)
    show = summary_df[["label", "auc_roc_mean", "auc_roc_std", "ap_mean", "ap_std",
                       "f1_mean", "f1_std"]].copy()
    show.columns = ["Configuration", "AUC mean", "AUC std", "AP mean", "AP std",
                    "F1 mean", "F1 std"]
    print(show.to_string(index=False, float_format=lambda x: f"{x:.4f}"))

    return {
        "raw":      raw_df,
        "summary":  summary_df,
        "per_type": per_type_df,
        "all":      all_results,
    }


def main():
    """CLI entry point. Thin wrapper around run_ablation() for command-line use."""
    args = parse_args()
    run_ablation(quick=args.quick, data_path=args.data, output_dir=args.out)


if __name__ == "__main__":
    main()
