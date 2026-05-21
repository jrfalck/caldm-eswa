"""
CALDM Ablation Study — ADBench Replication.

Companion to caldm_ablation_study_v3.py. Replicates the three most informative
architectural configurations from the BLS ablation (C, E, F) across every
.npz dataset found in a target folder, providing complementary evidence that
the architectural findings on BLS generalize to real-world tabular data.

Configurations evaluated (subset of the v3 ablation):
    C. Latent Diffusion (unconditional)        -- the closest non-context baseline
    E. CALDM-D (proposed)                      -- context conditions diffusion only
    F. CALDM-J                                 -- context conditions VAE and diffusion

Methodology
-----------
- Each .npz file in INPUT_DIR is loaded; each must contain `X` (features) and
  `y` (binary labels with 1=outlier). Other keys are ignored.
- Evaluation is TRANSDUCTIVE: each model is trained on the full dataset (no
  train/test split). This matches the standard unsupervised AD evaluation
  protocol used by ADBench itself, and is necessary for the many small
  datasets that cannot support a meaningful held-out set.
- DATA_FRACTION (default 1.0) controls what fraction of each dataset to use.
  Set to 0.5, 0.1, etc. for quick runs.
- Hyperparameters auto-scale to dataset shape (see scale_hyperparameters()).
- 3 seeds per (dataset, config) by default. Datasets failing on all seeds for
  any reason are skipped with a warning.

Outputs (under OUTPUT_DIR)
--------------------------
- adbench_per_run.csv      long format: one row per (dataset, config, seed)
- adbench_per_dataset.csv  wide format: one row per dataset (mean/std per config)
- adbench_summary.csv      aggregate across all datasets per config
- adbench_hyperparams.csv  hyperparameters used for each dataset (audit trail)
- adbench_results_table.docx  manuscript-ready table for Section 7.5
- figures/per_dataset_auroc.png  bar plot of AUROC per dataset per config

Run
---
    python caldm_ablation_adbench.py --input_dir path/to/adbench/folder
    python caldm_ablation_adbench.py --input_dir path/to/folder --quick
    python caldm_ablation_adbench.py --input_dir path/to/folder --data_fraction 0.5

Author: Juan
"""

import os
import time
import argparse
import random
import warnings
from dataclasses import dataclass, asdict

import numpy as np
import pandas as pd
import torch
import torch.nn as nn
import torch.optim as optim
import torch.nn.functional as F
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import (
    accuracy_score, precision_score, recall_score, f1_score,
    roc_auc_score, average_precision_score,
)
from scipy import stats
import matplotlib.pyplot as plt

warnings.filterwarnings("ignore", category=UserWarning, module="sklearn")
warnings.filterwarnings("ignore", message=".*enable_nested_tensor.*")
warnings.filterwarnings("ignore", message=".*Torch was not compiled with flash attention.*")

# =============================================================================
# 0. Configuration
# =============================================================================

INPUT_DIR = "data/adbench"
OUTPUT_DIR = "ablation_outputs_adbench"
DATA_FRACTION = 1.0
SEEDS = [42, 7, 123]

# Architectural defaults (will be auto-scaled per-dataset)
ATTENTION_HEADS_DEFAULT = 4
DROPOUT_RATE = 0.1
SALARY_WEIGHT = 1.0      # for ADBench (no privileged numerical feature), weight = 1
BOUND_CONTEXT = True

# Training schedule
VAE_EPOCHS = 15
VAE_BATCH_SIZE = 256
VAE_WARMUP = 5
VAE_MAX_BETA = 0.1 #1.0
DIFF_TIMESTEPS = 1000
DIFF_EPOCHS = 200
DIFF_BATCH_SIZE = 256
ENCODER_CHUNK_SIZE = 4096
LR = 1e-3
WEIGHT_DECAY = 1e-4

# Skip datasets that are too small to train meaningfully
MIN_ROWS = 100
MIN_FEATURES = 3

DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# =============================================================================
# Auto-scaling rule for per-dataset hyperparameters
# =============================================================================

def scale_hyperparameters(n_rows, n_features):
    """Compute per-dataset hyperparameters based on dataset shape.

    Rationale: ADBench datasets vary in size from ~100 rows to ~500k+ rows,
    and from 3 features to 500+ features. Using BLS hyperparameters
    (hidden=256, latent_cat=32, context=40) on a 100-row, 6-feature dataset
    massively over-parameterizes; using small hyperparameters on a 500-feature
    dataset under-parameterizes. The rule below scales sub-linearly with
    feature count and clamps to sensible bounds.

    Returns a dict reported in adbench_hyperparams.csv for full audit trail.
    """
    # Hidden dim: scales with sqrt(features), bounded [64, 256]
    hidden_dim = int(np.clip(64 * np.sqrt(max(n_features, 1) / 10), 64, 256))
    # round to nearest 32 for clean tensor shapes
    hidden_dim = max(64, (hidden_dim // 32) * 32)

    # Latent dim: roughly 1/4 of feature count, clamped, must be < n_features
    latent_dim = int(np.clip(max(4, n_features // 4), 4, 32))
    latent_dim = min(latent_dim, max(2, n_features - 1))

    # Context dim: similar shape to latent dim but capped at 40 (BLS default)
    context_dim = int(np.clip(max(4, n_features // 4), 4, 40))

    # Attention heads: must divide hidden_dim; pick largest power-of-2 divisor <= 4
    heads = ATTENTION_HEADS_DEFAULT
    while hidden_dim % heads != 0:
        heads -= 1
        if heads < 1:
            heads = 1
            break

    # Batch size: don't exceed sqrt(n_rows) for very small datasets
    batch_size = int(np.clip(VAE_BATCH_SIZE, 16, max(16, n_rows // 4)))
    batch_size = min(batch_size, n_rows)

    return {
        "n_rows": int(n_rows),
        "n_features": int(n_features),
        "hidden_dim": int(hidden_dim),
        "latent_dim": int(latent_dim),
        "context_dim": int(context_dim),
        "attention_heads": int(heads),
        "batch_size": int(batch_size),
    }


# =============================================================================
# 1. Reproducibility
# =============================================================================

def set_seed(seed):
    random.seed(seed)
    np.random.seed(seed)
    torch.manual_seed(seed)
    if torch.cuda.is_available():
        torch.cuda.manual_seed_all(seed)
    torch.backends.cudnn.deterministic = True
    torch.backends.cudnn.benchmark = False


# =============================================================================
# 2. Configurations
# =============================================================================

@dataclass
class AblationConfig:
    name: str
    label: str
    use_context: bool
    context_in_vae: bool
    context_in_diffusion: bool


CONFIGS = [
    AblationConfig("C_latent_diff_uncond", "Latent Diff (Uncond.)",
                   use_context=False, context_in_vae=False, context_in_diffusion=False),
    AblationConfig("E_caldm_d",            "CALDM-D (proposed)",
                   use_context=True,  context_in_vae=False, context_in_diffusion=True),
    AblationConfig("F_caldm_j",            "CALDM-J (jointly cond.)",
                   use_context=True,  context_in_vae=True,  context_in_diffusion=True),
]
PROPOSED_NAME = "E_caldm_d"


# =============================================================================
# 3. Models — single-headed VAE/diffusion (ADBench has no privileged feature)
# =============================================================================

class ContextEncoder(nn.Module):
    """Transformer-based context encoder. Identical to v3 except for default
    batch_first=False (preserves the praxis cross-sample attention behavior)
    and tanh-bounded output for joint-training stability."""

    def __init__(self, input_dim, embed_dim, hidden_dim, num_heads, dropout_rate):
        super().__init__()
        self.fc = nn.Sequential(
            nn.Linear(input_dim, hidden_dim), nn.ReLU(), nn.Dropout(dropout_rate),
        )
        encoder_layer = nn.TransformerEncoderLayer(
            d_model=hidden_dim, nhead=num_heads, dropout=dropout_rate
        )
        self.transformer = nn.TransformerEncoder(encoder_layer, num_layers=2)
        self.output_layer = nn.Linear(hidden_dim, embed_dim)

    def forward(self, x):
        h = self.fc(x).unsqueeze(1)
        h = self.transformer(h)
        out = self.output_layer(h.squeeze(1))
        return torch.tanh(out) if BOUND_CONTEXT else out


class FlexibleVAE(nn.Module):
    """Single-headed VAE for ADBench-style data (no salary/categorical split).
    Setting context_dim=0 cleanly disables conditioning."""

    def __init__(self, input_dim, latent_dim, context_dim, hidden_dim, dropout_rate=0.1):
        super().__init__()
        self.use_context = context_dim > 0
        self.context_dim = context_dim
        self.input_dim = input_dim
        self.latent_dim = latent_dim
        enc_in = input_dim + context_dim
        dec_in = latent_dim + context_dim

        self.encoder = self._mlp(enc_in, hidden_dim, latent_dim * 2, dropout_rate)
        self.decoder = self._mlp(dec_in, hidden_dim, input_dim, dropout_rate)

    @staticmethod
    def _mlp(in_dim, hid, out_dim, dropout):
        return nn.Sequential(
            nn.Linear(in_dim, hid), nn.LeakyReLU(), nn.Dropout(dropout),
            nn.Linear(hid, hid),    nn.LeakyReLU(), nn.Dropout(dropout),
            nn.Linear(hid, out_dim),
        )

    @staticmethod
    def reparameterize(mu, log_var):
        log_var = torch.clamp(log_var, min=-5, max=5)
        std = torch.exp(0.5 * log_var)
        return mu + torch.randn_like(std) * std

    def _maybe_cat(self, x, ctx):
        return torch.cat([x, ctx], dim=1) if self.use_context else x

    def encode(self, x, context=None):
        h = self.encoder(self._maybe_cat(x, context))
        mu, log_var = torch.chunk(h, 2, dim=1)
        log_var = torch.clamp(log_var, min=-5, max=5)
        return mu, log_var

    def decode(self, z, context=None):
        return self.decoder(self._maybe_cat(z, context))

    def forward(self, x, context=None):
        mu, log_var = self.encode(x, context)
        z = self.reparameterize(mu, log_var)
        return self.decode(z, context), mu, log_var


class FlexibleLatentDiffusion(nn.Module):
    """Single-headed latent DDPM. Linear beta schedule. context_dim=0 disables conditioning."""

    def __init__(self, latent_dim, context_dim, timesteps, hidden_dim, dropout_rate=0.1):
        super().__init__()
        self.timesteps = timesteps
        self.use_context = context_dim > 0
        self.context_dim = context_dim
        self.latent_dim = latent_dim

        in_dim = latent_dim + context_dim + 1

        beta = torch.linspace(0.0001, 0.02, timesteps)
        alpha = torch.clamp(1.0 - beta, min=1e-5, max=1.0)
        alpha_cum = torch.cumprod(alpha, dim=0)
        alpha_cum = torch.maximum(alpha_cum, torch.tensor(1e-5))
        self.register_buffer("beta", beta)
        self.register_buffer("alpha", alpha)
        self.register_buffer("alpha_cum", alpha_cum)

        self.network = nn.Sequential(
            nn.Linear(in_dim, hidden_dim), nn.ReLU(), nn.Dropout(dropout_rate),
            nn.Linear(hidden_dim, hidden_dim), nn.ReLU(), nn.Dropout(dropout_rate),
            nn.Linear(hidden_dim, latent_dim),
        )

    def forward_diffusion(self, z0, t):
        alpha_t = self.alpha_cum[t].unsqueeze(1)
        n = torch.randn_like(z0)
        z_t = torch.sqrt(alpha_t) * z0 + torch.sqrt(1 - alpha_t) * n
        return z_t, n

    def reverse_denoising(self, z_t, context, t):
        t_norm = t.view(-1, 1).float() / self.timesteps
        if self.use_context and context is not None:
            x_in = torch.cat([z_t, context, t_norm], dim=1)
        else:
            x_in = torch.cat([z_t, t_norm], dim=1)
        return 0.5 * self.network(x_in).tanh()

    @torch.no_grad()
    def sample(self, batch_size, context, steps, device):
        z = torch.randn(batch_size, self.latent_dim, device=device)
        for t in reversed(range(steps)):
            z = z / (torch.norm(z, p=2, dim=1, keepdim=True) + 1e-5)
            t_long = torch.full((batch_size,), t, device=device, dtype=torch.long)
            pred = self.reverse_denoising(z, context, t_long)
            alpha_t = torch.clamp(self.alpha_cum[t].view(-1, 1), min=1e-3, max=1.0)
            z = (z - torch.sqrt(1 - alpha_t) * pred) / (torch.sqrt(alpha_t) + 1e-5)
            if t > 0:
                z = z + 0.05 * torch.sqrt(self.beta[t]) * torch.randn_like(z)
        return z


# =============================================================================
# 4. Training and scoring
# =============================================================================

def beta_scheduler(epoch, warmup_epochs, max_beta):
    if epoch < warmup_epochs:
        return max(0.005, (epoch / warmup_epochs) ** 2 * max_beta)
    return max_beta


def train_vae(vae, context_encoder, X, epochs, batch_size, device, verbose=False):
    params = list(vae.parameters())
    if context_encoder is not None:
        params += list(context_encoder.parameters())
    opt = optim.Adam(params, lr=LR, weight_decay=WEIGHT_DECAY)

    X_aug = X + torch.randn_like(X) * 0.01
    loader = torch.utils.data.DataLoader(
        torch.utils.data.TensorDataset(X_aug),
        batch_size=batch_size, shuffle=True
    )

    for epoch in range(epochs):
        vae.train()
        if context_encoder is not None:
            context_encoder.train()
        beta = beta_scheduler(epoch, VAE_WARMUP, VAE_MAX_BETA)
        for (xb,) in loader:
            xb = xb.to(device)
            opt.zero_grad()
            ctx = context_encoder(xb) if context_encoder is not None else None
            x_rec, mu, log_var = vae(xb, ctx)
            kl = -0.5 * torch.sum(1 + log_var - mu.pow(2) - log_var.exp()) / xb.size(0)
            loss = F.mse_loss(x_rec, xb) + beta * kl
            if not torch.isfinite(loss):
                continue
            loss.backward()
            grad_finite = all((p.grad is None) or torch.isfinite(p.grad).all() for p in params)
            if not grad_finite:
                opt.zero_grad()
                continue
            torch.nn.utils.clip_grad_norm_(params, max_norm=1.0)
            opt.step()

    # NaN safety check
    for name, p in (list(vae.named_parameters())
                    + (list(context_encoder.named_parameters()) if context_encoder is not None else [])):
        if not torch.isfinite(p).all():
            raise RuntimeError(f"VAE/encoder produced non-finite weights in '{name}'.")

    vae.eval()
    if context_encoder is not None:
        context_encoder.eval()


def train_latent_diffusion(diff, vae, context_encoder, config, X, epochs, device,
                           batch_size, verbose=False):
    train_encoder = (
        config.context_in_diffusion
        and not config.context_in_vae
        and context_encoder is not None
    )

    params = list(diff.parameters())
    if train_encoder:
        params += list(context_encoder.parameters())
    opt = optim.Adam(params, lr=LR)

    X_d = X.to(device)
    n = X_d.shape[0]
    vae.eval()

    # Precompute frozen latents
    z_chunks = []
    ctx_diff_chunks = []
    use_precomputed_ctx = (
        config.context_in_diffusion and context_encoder is not None and not train_encoder
    )

    with torch.no_grad():
        for s in range(0, n, batch_size):
            xb = X_d[s:s + batch_size]
            ctx_vae = (context_encoder(xb)
                       if (config.context_in_vae and context_encoder is not None) else None)
            mu, lv = vae.encode(xb, ctx_vae)
            z_chunks.append(vae.reparameterize(mu, lv))
            if use_precomputed_ctx:
                ctx_diff_chunks.append(context_encoder(xb))

    z_train = torch.cat(z_chunks, dim=0)
    ctx_diff_pre = torch.cat(ctx_diff_chunks, dim=0) if ctx_diff_chunks else None

    for epoch in range(epochs):
        diff.train()
        if train_encoder:
            context_encoder.train()

        perm = torch.randperm(n, device=device)
        for s in range(0, n, batch_size):
            idx = perm[s:s + batch_size]
            opt.zero_grad()
            zb = z_train[idx]

            if config.context_in_diffusion and context_encoder is not None:
                if train_encoder:
                    ctx_diff = context_encoder(X_d[idx])
                else:
                    ctx_diff = ctx_diff_pre[idx]
            else:
                ctx_diff = None

            t = torch.randint(0, diff.timesteps, (zb.shape[0],), device=device)
            z_t, noise = diff.forward_diffusion(zb, t)
            pred = diff.reverse_denoising(z_t, ctx_diff, t)
            loss = F.mse_loss(pred, noise)
            if not torch.isfinite(loss):
                continue
            loss.backward()
            torch.nn.utils.clip_grad_norm_(params, max_norm=1.0)
            opt.step()

    diff.eval()
    if context_encoder is not None:
        context_encoder.eval()


@torch.no_grad()
def score_latent_pipeline(vae, diff, context_encoder, config, X, device,
                          chunk_size=ENCODER_CHUNK_SIZE):
    X = X.to(device)
    n = X.shape[0]
    parts = []
    for s in range(0, n, chunk_size):
        xb = X[s:s + chunk_size]
        ctx_vae = (context_encoder(xb)
                   if (context_encoder is not None and config.context_in_vae) else None)
        ctx_diff = (context_encoder(xb)
                    if (context_encoder is not None and config.context_in_diffusion) else None)
        z = diff.sample(batch_size=xb.shape[0], context=ctx_diff,
                        steps=diff.timesteps, device=device)
        x_rec = vae.decode(z, ctx_vae)
        loss = torch.mean((xb - x_rec) ** 2, dim=1)
        parts.append(torch.log1p(loss))
    return torch.cat(parts).cpu().numpy()


# =============================================================================
# 5. Per-dataset experiment runner
# =============================================================================

def run_single_experiment(config, seed, X_t, y, hp, device, verbose=False):
    set_seed(seed)
    t0 = time.time()

    ctx_dim_vae = hp["context_dim"] if config.context_in_vae else 0
    ctx_dim_diff = hp["context_dim"] if config.context_in_diffusion else 0

    context_encoder = None
    if config.use_context:
        context_encoder = ContextEncoder(
            hp["n_features"], embed_dim=hp["context_dim"],
            hidden_dim=hp["hidden_dim"], num_heads=hp["attention_heads"],
            dropout_rate=DROPOUT_RATE
        ).to(device)

    vae = FlexibleVAE(
        hp["n_features"], hp["latent_dim"], ctx_dim_vae, hp["hidden_dim"], DROPOUT_RATE
    ).to(device)
    ctx_for_vae_train = context_encoder if config.context_in_vae else None
    train_vae(vae, ctx_for_vae_train, X_t, VAE_EPOCHS, hp["batch_size"], device, verbose)

    diff = FlexibleLatentDiffusion(
        hp["latent_dim"], ctx_dim_diff, DIFF_TIMESTEPS, hp["hidden_dim"], DROPOUT_RATE
    ).to(device)
    train_latent_diffusion(diff, vae, context_encoder, config, X_t, DIFF_EPOCHS,
                           device, hp["batch_size"], verbose)

    scores = score_latent_pipeline(vae, diff, context_encoder, config, X_t, device)
    train_time = time.time() - t0

    if not np.all(np.isfinite(scores)):
        raise ValueError(f"{int((~np.isfinite(scores)).sum())} non-finite scores")

    metrics = compute_metrics(scores, y)
    metrics["train_time_sec"] = train_time
    return metrics


def compute_metrics(scores, y_true):
    outlier_rate = float(y_true.mean())
    threshold = float(np.percentile(scores, 100 * (1 - outlier_rate)))
    y_pred = (scores >= threshold).astype(int)
    return {
        "auc_roc":   roc_auc_score(y_true, scores),
        "ap":        average_precision_score(y_true, scores),
        "f1":        f1_score(y_true, y_pred, zero_division=0),
        "precision": precision_score(y_true, y_pred, zero_division=0),
        "recall":    recall_score(y_true, y_pred, zero_division=0),
        "threshold": threshold,
        "outlier_rate": outlier_rate,
    }


# =============================================================================
# 6. Data loading
# =============================================================================

def load_adbench_dataset(filepath, data_fraction):
    """Load a single .npz file, expect X and y. Subsample if data_fraction < 1.
    Returns (X_tensor, y_array, metadata_dict) or None if dataset is invalid."""
    try:
        data = np.load(filepath, allow_pickle=True)
    except Exception as e:
        return None, f"Failed to load: {e}"

    if "X" not in data or "y" not in data:
        return None, "Missing X or y key"

    X = np.array(data["X"], dtype=np.float32)
    y = np.array(data["y"]).astype(int).flatten()

    if X.ndim != 2:
        return None, f"X has unexpected shape {X.shape}"
    if X.shape[0] != y.shape[0]:
        return None, f"X.shape[0]={X.shape[0]} != y.shape[0]={y.shape[0]}"
    if X.shape[0] < MIN_ROWS:
        return None, f"Too few rows ({X.shape[0]} < {MIN_ROWS})"
    if X.shape[1] < MIN_FEATURES:
        return None, f"Too few features ({X.shape[1]} < {MIN_FEATURES})"
    if len(np.unique(y)) < 2:
        return None, "Only one class present in y"

    # Subsample (stratified) if requested
    if data_fraction < 1.0:
        rng = np.random.RandomState(42)
        n_keep = int(X.shape[0] * data_fraction)
        # stratified subsample to preserve outlier rate
        n_pos = int(round(y.mean() * n_keep))
        n_neg = n_keep - n_pos
        pos_idx = np.where(y == 1)[0]
        neg_idx = np.where(y == 0)[0]
        if len(pos_idx) < n_pos: n_pos = len(pos_idx)
        if len(neg_idx) < n_neg: n_neg = len(neg_idx)
        keep_idx = np.concatenate([
            rng.choice(pos_idx, n_pos, replace=False),
            rng.choice(neg_idx, n_neg, replace=False),
        ])
        rng.shuffle(keep_idx)
        X = X[keep_idx]
        y = y[keep_idx]

    # Check post-subsample validity
    if len(np.unique(y)) < 2:
        return None, "Only one class after subsampling"

    # Standardize features
    scaler = StandardScaler()
    X_std = scaler.fit_transform(X)
    X_t = torch.tensor(X_std, dtype=torch.float32)

    meta = {
        "n_rows": int(X.shape[0]),
        "n_features": int(X.shape[1]),
        "outlier_rate": float(y.mean()),
        "n_outliers": int(y.sum()),
    }
    return (X_t, y, meta), None


# =============================================================================
# 7. Output writers
# =============================================================================

def write_per_dataset_table(per_dataset_df, output_path):
    """Build the manuscript-ready DOCX table. Uses landscape orientation
    because the per-dataset table is naturally wide (10 columns)."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[warn] python-docx not installed; skipping .docx output")
        return

    doc = Document()
    # Landscape orientation for the wide table
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading("CALDM Ablation \u2014 ADBench Replication", level=1)
    doc.add_paragraph(
        f"Per-dataset AUROC and AP for the three ablation configurations across "
        f"{len(per_dataset_df)} ADBench datasets. Each cell reports mean \u00B1 std "
        f"across {len(SEEDS)} random seeds. Bold marks the best config per metric per dataset."
    )

    # Build table: Dataset | rows | features | outlier % | C-AUC | E-AUC | F-AUC | C-AP | E-AP | F-AP
    headers = ["Dataset", "Rows", "Features", "Outlier %",
               "AUROC C", "AUROC E (CALDM-D)", "AUROC F (CALDM-J)",
               "AP C", "AP E (CALDM-D)", "AP F (CALDM-J)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)

    for _, row in per_dataset_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["dataset"])
        cells[1].text = str(int(row["n_rows"]))
        cells[2].text = str(int(row["n_features"]))
        cells[3].text = f"{row['outlier_rate']*100:.2f}%"

        # Identify best per metric across configs
        auc_vals = [row.get("C_latent_diff_uncond_auc_roc_mean", np.nan),
                    row.get("E_caldm_d_auc_roc_mean", np.nan),
                    row.get("F_caldm_j_auc_roc_mean", np.nan)]
        ap_vals = [row.get("C_latent_diff_uncond_ap_mean", np.nan),
                   row.get("E_caldm_d_ap_mean", np.nan),
                   row.get("F_caldm_j_ap_mean", np.nan)]
        best_auc = int(np.nanargmax(auc_vals)) if not all(np.isnan(auc_vals)) else -1
        best_ap = int(np.nanargmax(ap_vals)) if not all(np.isnan(ap_vals)) else -1

        for cfg_i, cfg in enumerate(["C_latent_diff_uncond", "E_caldm_d", "F_caldm_j"]):
            auc_m = row.get(f"{cfg}_auc_roc_mean", np.nan)
            auc_s = row.get(f"{cfg}_auc_roc_std", np.nan)
            text = f"{auc_m:.3f}\u00B1{auc_s:.3f}" if not np.isnan(auc_m) else "\u2014"
            cells[4 + cfg_i].text = text
            if cfg_i == best_auc:
                for run in cells[4 + cfg_i].paragraphs[0].runs:
                    run.bold = True
            for run in cells[4 + cfg_i].paragraphs[0].runs:
                run.font.size = Pt(8)

        for cfg_i, cfg in enumerate(["C_latent_diff_uncond", "E_caldm_d", "F_caldm_j"]):
            ap_m = row.get(f"{cfg}_ap_mean", np.nan)
            ap_s = row.get(f"{cfg}_ap_std", np.nan)
            text = f"{ap_m:.3f}\u00B1{ap_s:.3f}" if not np.isnan(ap_m) else "\u2014"
            cells[7 + cfg_i].text = text
            if cfg_i == best_ap:
                for run in cells[7 + cfg_i].paragraphs[0].runs:
                    run.bold = True
            for run in cells[7 + cfg_i].paragraphs[0].runs:
                run.font.size = Pt(8)

    doc.save(output_path)
    print(f"[ok] Saved {output_path}")


def make_per_dataset_plot(per_dataset_df, output_path):
    n = len(per_dataset_df)
    fig, ax = plt.subplots(figsize=(max(10, n * 0.35), 5))
    x = np.arange(n)
    width = 0.27
    cfg_labels = {
        "C_latent_diff_uncond": ("C: Latent Diff (Uncond.)", "#3b7dd8"),
        "E_caldm_d":            ("E: CALDM-D (proposed)",    "#d8633b"),
        "F_caldm_j":            ("F: CALDM-J",                "#7f7f7f"),
    }
    for i, (cfg, (label, color)) in enumerate(cfg_labels.items()):
        means = per_dataset_df[f"{cfg}_auc_roc_mean"].values
        stds  = per_dataset_df[f"{cfg}_auc_roc_std"].values
        ax.bar(x + (i - 1) * width, means, width, yerr=stds, capsize=2,
               label=label, color=color, edgecolor="black", linewidth=0.3)
    ax.set_xticks(x)
    ax.set_xticklabels(per_dataset_df["dataset"].values, rotation=45, ha="right", fontsize=8)
    ax.set_ylabel("AUROC")
    ax.set_title(f"Per-Dataset AUROC across {n} ADBench Datasets (mean \u00B1 std, {len(SEEDS)} seeds)")
    ax.axhline(0.5, color="black", linestyle="--", alpha=0.3, linewidth=0.8)
    ax.legend(loc="lower right")
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()


# =============================================================================
# 8. Main
# =============================================================================

def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--input_dir", type=str, default=INPUT_DIR,
                   help="Folder containing ADBench .npz files")
    p.add_argument("--out", type=str, default=OUTPUT_DIR, help="Output directory")
    p.add_argument("--data_fraction", type=float, default=DATA_FRACTION,
                   help="Fraction of each dataset to use (default 1.0 = full)")
    p.add_argument("--seeds", type=int, nargs="+", default=None,
                   help="Random seeds (default: 3 seeds)")
    p.add_argument("--quick", action="store_true",
                   help="Quick mode: 1 seed, fewer epochs")
    p.add_argument("--verbose", action="store_true")
    return p.parse_args()


def run_adbench_ablation(input_dir=None, output_dir=None, data_fraction=None,
                        seeds=None, vae_epochs=None, diff_epochs=None,
                        quick=False, verbose=False):
    """Programmatic entry point. Use this from Jupyter or another script.

    All parameters are optional; if not given, the module-level defaults are used.

    Example (Jupyter):
        from caldm_ablation_adbench import run_adbench_ablation

        # Default: 3 seeds, full datasets
        results = run_adbench_ablation(input_dir="path/to/adbench/folder")

        # Quick smoke test
        results = run_adbench_ablation(input_dir="path/to/adbench/folder", quick=True)

        # Use half of each dataset to speed up
        results = run_adbench_ablation(input_dir="path/to/adbench/folder",
                                       data_fraction=0.5)

        # Inspect after running
        results["per_dataset"]   # one row per dataset, all configs side by side
        results["per_run"]       # one row per (dataset, config, seed)
        results["summary"]       # aggregate per config across all datasets
        results["hyperparams"]   # what hyperparameters were used per dataset
        results["skipped"]       # datasets that couldn't run

    Returns
    -------
    dict with keys: 'per_run', 'per_dataset', 'summary', 'hyperparams', 'skipped'
    """
    global VAE_EPOCHS, DIFF_EPOCHS, SEEDS, OUTPUT_DIR, INPUT_DIR, DATA_FRACTION

    if quick:
        if seeds is None:           seeds = [42]
        if vae_epochs is None:      vae_epochs = 5
        if diff_epochs is None:     diff_epochs = 5
        print("[setup] QUICK mode: 1 seed, 5 epochs each")

    # Apply explicit overrides last so they win
    if input_dir is not None:        INPUT_DIR = input_dir
    if output_dir is not None:       OUTPUT_DIR = output_dir
    if data_fraction is not None:    DATA_FRACTION = data_fraction
    if seeds is not None:            SEEDS = list(seeds)
    if vae_epochs is not None:       VAE_EPOCHS = vae_epochs
    if diff_epochs is not None:      DIFF_EPOCHS = diff_epochs

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(os.path.join(OUTPUT_DIR, "figures"), exist_ok=True)

    if not os.path.isdir(INPUT_DIR):
        raise FileNotFoundError(f"Input directory not found: {INPUT_DIR}")

    npz_files = sorted([f for f in os.listdir(INPUT_DIR) if f.endswith(".npz")])
    if not npz_files:
        raise FileNotFoundError(f"No .npz files found in {INPUT_DIR}")

    print(f"[setup] device={DEVICE}  seeds={SEEDS}  data_fraction={DATA_FRACTION}")
    print(f"[setup] found {len(npz_files)} .npz files in {INPUT_DIR}")

    per_run_records = []
    hp_records = []
    skipped = []

    overall_t0 = time.time()
    for fi, fname in enumerate(npz_files):
        dataset_name = os.path.splitext(fname)[0]
        filepath = os.path.join(INPUT_DIR, fname)

        loaded, err = load_adbench_dataset(filepath, DATA_FRACTION)
        if loaded is None:
            print(f"\n[skip] {dataset_name}: {err}")
            skipped.append({"dataset": dataset_name, "reason": err})
            continue

        X_t, y, meta = loaded
        hp = scale_hyperparameters(meta["n_rows"], meta["n_features"])
        hp_record = {"dataset": dataset_name, **hp,
                     "outlier_rate": meta["outlier_rate"],
                     "n_outliers": meta["n_outliers"]}
        hp_records.append(hp_record)

        print(f"\n[{fi+1}/{len(npz_files)}] {dataset_name}  "
              f"({meta['n_rows']} rows, {meta['n_features']} features, "
              f"{meta['outlier_rate']*100:.2f}% outliers)")
        print(f"  hp: hidden={hp['hidden_dim']} latent={hp['latent_dim']} "
              f"context={hp['context_dim']} heads={hp['attention_heads']} "
              f"batch={hp['batch_size']}")

        for cfg in CONFIGS:
            for seed in SEEDS:
                try:
                    metrics = run_single_experiment(cfg, seed, X_t, y, hp, DEVICE, verbose)
                    per_run_records.append({
                        "dataset": dataset_name,
                        "config": cfg.name,
                        "config_label": cfg.label,
                        "seed": seed,
                        **meta,
                        **metrics,
                    })
                    print(f"  {cfg.label} seed={seed}: "
                          f"AUC={metrics['auc_roc']:.3f} AP={metrics['ap']:.3f} "
                          f"({metrics['train_time_sec']:.1f}s)")
                except Exception as e:
                    print(f"  [error] {cfg.label} seed={seed}: {e}")
                    per_run_records.append({
                        "dataset": dataset_name,
                        "config": cfg.name,
                        "config_label": cfg.label,
                        "seed": seed,
                        **meta,
                        "auc_roc": np.nan, "ap": np.nan, "f1": np.nan,
                        "precision": np.nan, "recall": np.nan,
                        "error": str(e),
                    })

    total_time = time.time() - overall_t0
    print(f"\n[done] Total run time: {total_time/60:.1f} minutes")

    # ---- Long-format CSV
    per_run_df = pd.DataFrame(per_run_records)
    per_run_df.to_csv(os.path.join(OUTPUT_DIR, "adbench_per_run.csv"), index=False)
    print(f"[ok] Saved adbench_per_run.csv ({len(per_run_df)} rows)")

    # ---- Hyperparameters audit trail
    hp_df = pd.DataFrame(hp_records)
    hp_df.to_csv(os.path.join(OUTPUT_DIR, "adbench_hyperparams.csv"), index=False)
    print(f"[ok] Saved adbench_hyperparams.csv")

    skipped_df = pd.DataFrame(skipped) if skipped else pd.DataFrame()
    if skipped:
        skipped_df.to_csv(os.path.join(OUTPUT_DIR, "adbench_skipped.csv"), index=False)
        print(f"[ok] Saved adbench_skipped.csv ({len(skipped)} datasets)")

    # ---- Per-dataset wide format
    if per_run_df.empty:
        print("[warn] No successful runs; skipping aggregate outputs")
        return {
            "per_run": per_run_df,
            "per_dataset": pd.DataFrame(),
            "summary": pd.DataFrame(),
            "hyperparams": hp_df,
            "skipped": skipped_df,
        }

    metric_cols = ["auc_roc", "ap", "f1", "precision", "recall"]
    grouped = (per_run_df.groupby(["dataset", "config"])[metric_cols]
               .agg(["mean", "std"]).reset_index())
    grouped.columns = ["dataset", "config"] + [f"{m}_{s}" for m, s in grouped.columns[2:]]

    wide = grouped.pivot(index="dataset", columns="config")
    wide.columns = [f"{cfg}_{m}" for m, cfg in wide.columns]
    wide = wide.reset_index()

    meta_df = (per_run_df.groupby("dataset")[["n_rows", "n_features", "outlier_rate"]]
               .first().reset_index())
    per_dataset_df = wide.merge(meta_df, on="dataset")
    per_dataset_df.to_csv(os.path.join(OUTPUT_DIR, "adbench_per_dataset.csv"), index=False)
    print(f"[ok] Saved adbench_per_dataset.csv ({len(per_dataset_df)} datasets)")

    # ---- Aggregate summary across all datasets
    summary_records = []
    for cfg in CONFIGS:
        cfg_df = per_run_df[per_run_df["config"] == cfg.name]
        if cfg_df.empty:
            continue

        per_ds_means = (cfg_df.groupby("dataset")[["auc_roc", "ap"]]
                        .mean().reset_index())
        summary_records.append({
            "config": cfg.name,
            "label": cfg.label,
            "n_datasets": int(per_ds_means.shape[0]),
            "auc_roc_mean": float(per_ds_means["auc_roc"].mean()),
            "auc_roc_std": float(per_ds_means["auc_roc"].std()),
            "ap_mean": float(per_ds_means["ap"].mean()),
            "ap_std": float(per_ds_means["ap"].std()),
        })

    win_auc = {c.name: 0 for c in CONFIGS}
    win_ap = {c.name: 0 for c in CONFIGS}
    for ds in per_dataset_df["dataset"].unique():
        auc_means = {c.name: per_dataset_df.loc[per_dataset_df["dataset"] == ds,
                                                f"{c.name}_auc_roc_mean"].iloc[0]
                     for c in CONFIGS}
        ap_means = {c.name: per_dataset_df.loc[per_dataset_df["dataset"] == ds,
                                               f"{c.name}_ap_mean"].iloc[0]
                    for c in CONFIGS}
        if not all(np.isnan(list(auc_means.values()))):
            best_auc_cfg = max(auc_means, key=lambda k: auc_means[k] if not np.isnan(auc_means[k]) else -np.inf)
            win_auc[best_auc_cfg] += 1
        if not all(np.isnan(list(ap_means.values()))):
            best_ap_cfg = max(ap_means, key=lambda k: ap_means[k] if not np.isnan(ap_means[k]) else -np.inf)
            win_ap[best_ap_cfg] += 1

    for rec in summary_records:
        rec["wins_auc_roc"] = win_auc[rec["config"]]
        rec["wins_ap"] = win_ap[rec["config"]]

    proposed_per_ds = (per_run_df[per_run_df["config"] == PROPOSED_NAME]
                       .groupby("dataset")[["auc_roc", "ap"]].mean()
                       .reset_index().sort_values("dataset"))
    for cfg in CONFIGS:
        if cfg.name == PROPOSED_NAME:
            continue
        cfg_per_ds = (per_run_df[per_run_df["config"] == cfg.name]
                      .groupby("dataset")[["auc_roc", "ap"]].mean()
                      .reset_index().sort_values("dataset"))
        merged = proposed_per_ds.merge(cfg_per_ds, on="dataset",
                                       suffixes=("_proposed", "_other"))
        if len(merged) >= 2:
            try:
                _, p_auc = stats.ttest_rel(merged["auc_roc_proposed"], merged["auc_roc_other"])
                _, p_ap = stats.ttest_rel(merged["ap_proposed"], merged["ap_other"])
            except Exception:
                p_auc, p_ap = np.nan, np.nan
        else:
            p_auc, p_ap = np.nan, np.nan
        for rec in summary_records:
            if rec["config"] == cfg.name:
                rec["p_auc_vs_proposed"] = p_auc
                rec["p_ap_vs_proposed"] = p_ap

    summary_df = pd.DataFrame(summary_records)
    summary_df.to_csv(os.path.join(OUTPUT_DIR, "adbench_summary.csv"), index=False)
    print(f"[ok] Saved adbench_summary.csv")

    write_per_dataset_table(per_dataset_df,
                            os.path.join(OUTPUT_DIR, "adbench_results_table.docx"))

    make_per_dataset_plot(per_dataset_df,
                          os.path.join(OUTPUT_DIR, "figures", "per_dataset_auroc.png"))
    print(f"[ok] Saved figures/per_dataset_auroc.png")

    print("\n" + "=" * 80)
    print(f"AGGREGATE SUMMARY ACROSS {len(per_dataset_df)} DATASETS")
    print("=" * 80)
    show = summary_df[["label", "auc_roc_mean", "auc_roc_std", "ap_mean", "ap_std",
                       "wins_auc_roc", "wins_ap"]].copy()
    show.columns = ["Configuration", "AUROC mean", "AUROC std", "AP mean", "AP std",
                    "Wins AUROC", "Wins AP"]
    print(show.to_string(index=False, float_format=lambda x: f"{x:.4f}"))
    if "p_auc_vs_proposed" in summary_df.columns:
        print("\nPaired t-tests vs. CALDM-D (across datasets):")
        for _, row in summary_df.iterrows():
            if row["config"] == PROPOSED_NAME:
                continue
            print(f"  {row['label']:<30s}  "
                  f"p_AUROC = {row.get('p_auc_vs_proposed', np.nan):.4f}  "
                  f"p_AP = {row.get('p_ap_vs_proposed', np.nan):.4f}")

    return {
        "per_run":     per_run_df,
        "per_dataset": per_dataset_df,
        "summary":     summary_df,
        "hyperparams": hp_df,
        "skipped":     skipped_df,
    }


def main():
    """CLI entry point. Thin wrapper around run_adbench_ablation()."""
    args = parse_args()
    run_adbench_ablation(
        input_dir=args.input_dir,
        output_dir=args.out,
        data_fraction=args.data_fraction,
        seeds=args.seeds,
        quick=args.quick,
        verbose=args.verbose,
    )


if __name__ == "__main__":
    main()
